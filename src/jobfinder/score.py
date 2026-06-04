"""Résumé extraction and scoring (LLD §6).

This module owns the semantic-matching half of the pipeline:

* :func:`extract_resume` reads the user's full résumé from
  ``config/resume.{pdf,docx,txt,md}`` into plain text (T14, LLD §6.5).
* :func:`build_profile_vector` and :func:`embed_job` turn text into the
  L2-normalized embeddings the scorer compares (T15, LLD §6.1–§6.3). The final
  weighted ``score_job`` lands with T16.

The heavyweight ``sentence-transformers``/torch import is kept lazy inside
:func:`load_model` so importing this module (e.g. to extract a résumé) stays
cheap and never loads torch until an embedding is actually requested.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING, Protocol

import numpy as np

from .models import LocationBucket, ScoreBreakdown

if TYPE_CHECKING:
    from datetime import datetime

    from numpy.typing import NDArray

    from .models import Job
    from .settings import Profile, Weights

# Supported résumé formats, dispatched on the lowercased file suffix (LLD §6.5).
_TEXT_SUFFIXES = frozenset({".txt", ".md"})
_PDF_SUFFIX = ".pdf"
_DOCX_SUFFIX = ".docx"
_SUPPORTED_SUFFIXES = _TEXT_SUFFIXES | {_PDF_SUFFIX, _DOCX_SUFFIX}

# Profile chunking: the resume + targeting block is split into word-windows that
# stay under the model's input limit so the tail is never truncated (LLD §6.2).
# all-MiniLM-L6-v2 truncates at 256 tokens; ~180 English words ≈ 256 tokens, so
# a conservative word window keeps every chunk inside the limit without needing
# the tokenizer at chunk time (keeps chunking pure + offline-testable).
_PROFILE_CHUNK_MAX_WORDS = 180  # ≈256 tokens, LLD §6.2

# A single job is embedded from title+description in one pass (LLD §6.3); the
# char cap bounds tokenizer work on pathological descriptions. The model still
# truncates to its own token limit — this only stops us tokenizing megabytes.
_JOB_CHAR_CAP = 5000  # LLD §6.3 CHAR_CAP


def extract_resume(path: str | Path) -> str:
    """Extract the full plain text of a résumé file (LLD §6.5).

    Dispatches on the file extension: ``.pdf`` via pypdf (falling back to
    pdfplumber when pypdf yields no text), ``.docx`` via python-docx
    (paragraphs and tables), ``.txt``/``.md`` read directly as UTF-8.

    Raises:
        FileNotFoundError: the résumé file does not exist.
        ValueError: the file extension is not a supported résumé format.
    """
    resume_path = Path(path)
    if not resume_path.exists():
        raise FileNotFoundError(f"résumé file not found: {resume_path}")

    suffix = resume_path.suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return resume_path.read_text(encoding="utf-8")
    if suffix == _DOCX_SUFFIX:
        return _extract_docx(resume_path)
    if suffix == _PDF_SUFFIX:
        return _extract_pdf(resume_path)

    supported = ", ".join(sorted(_SUPPORTED_SUFFIXES))
    raise ValueError(
        f"unsupported résumé format {suffix!r} for {resume_path}; supported: {supported}"
    )


def _extract_pdf(path: Path) -> str:
    """Extract PDF text via pypdf, falling back to pdfplumber if empty (LLD §6.5)."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if text.strip():
        return text
    # pypdf produced nothing usable (empty/garbled layout) — try the heavier,
    # more layout-tolerant pdfplumber extractor before giving up.
    return _extract_pdf_pdfplumber(path)


def _extract_pdf_pdfplumber(path: Path) -> str:
    """Fallback PDF extraction via pdfplumber for layouts pypdf can't read."""
    import pdfplumber

    with pdfplumber.open(str(path)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _extract_docx(path: Path) -> str:
    """Extract docx paragraphs and table cells in document order (LLD §6.5)."""
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class Encoder(Protocol):
    """Minimal embedding-model interface (SentenceTransformer-compatible).

    Declared as a Protocol so the pipeline can pass the real model while tests
    inject a deterministic fake encoder, keeping the chunk/pool math offline.
    """

    def encode(self, sentences: str | list[str], *, normalize_embeddings: bool = ...): ...


# SentenceTransformer instances keyed by model name; loading is expensive, so a
# poll (and the test session) reuses one instance across calls (LLD §6.1).
_MODEL_CACHE: dict[str, Encoder] = {}


def load_model(name: str) -> Encoder:
    """Return a cached :class:`SentenceTransformer` for ``name`` (LLD §6.1).

    The first call constructs and caches the model (downloaded once into the HF
    cache, then offline); later calls for the same name reuse it. The
    ``sentence_transformers``/torch import is deferred to here so callers that
    only need résumé extraction never pay it.
    """
    cached = _MODEL_CACHE.get(name)
    if cached is not None:
        return cached
    from sentence_transformers import SentenceTransformer

    model = SentenceTransformer(name)
    _MODEL_CACHE[name] = model
    return model


def render_targeting(profile: Profile) -> str:
    """Render the structured targeting block prepended to the résumé (LLD §6.2).

    Putting role/skills/seniority first lets the user's stated priorities steer
    the profile vector rather than being diluted by the full résumé text.
    """
    return (
        f"Target role: {', '.join(profile.role_keywords)}.\n"
        f"Must-have skills: {', '.join(profile.must_have_skills)}.\n"
        f"Seniority: {', '.join(profile.seniority_include)}."
    )


def _chunk_text(text: str, max_words: int) -> list[str]:
    """Split ``text`` into word-windows of at most ``max_words`` (pure helper).

    Always returns at least one chunk so the caller gets a usable embedding even
    for empty input.
    """
    words = text.split()
    if not words:
        return [text]
    return [" ".join(words[i : i + max_words]) for i in range(0, len(words), max_words)]


def _l2_normalize(vec: NDArray[np.float32]) -> NDArray[np.float32]:
    """Return ``vec`` scaled to unit length (zero vector returned unchanged)."""
    norm = float(np.linalg.norm(vec))
    if norm == 0.0:
        return vec
    return (vec / norm).astype(np.float32)


def build_profile_vector(
    profile: Profile, resume_text: str, *, model: Encoder
) -> NDArray[np.float32]:
    """Build the single L2-normalized profile vector for a poll (LLD §6.2).

    The targeting block is prepended to the full résumé, the combined text is
    chunked so the tail survives the model's input limit, each chunk is encoded
    (already L2-normalized), the chunk vectors are mean-pooled, and the pooled
    vector is re-normalized to unit length.
    """
    profile_text = f"{render_targeting(profile)}\n\n{resume_text}"
    chunks = _chunk_text(profile_text, _PROFILE_CHUNK_MAX_WORDS)
    chunk_vecs = np.asarray(model.encode(chunks, normalize_embeddings=True), dtype=np.float32)
    pooled = chunk_vecs.mean(axis=0)
    return _l2_normalize(pooled)


def embed_job(job: Job, *, model: Encoder) -> NDArray[np.float32]:
    """Return the L2-normalized embedding of a job's title+description (LLD §6.3)."""
    job_text = f"{job.title}\n{job.description}"[:_JOB_CHAR_CAP]
    vec = np.asarray(model.encode(job_text, normalize_embeddings=True), dtype=np.float32)
    return _l2_normalize(vec)


# Location bonus by bucket (LLD §6.3): remote > vancouver > toronto > other_canada,
# anything outside Canada/remote scores zero (those are filtered out anyway).
_LOCATION_BONUS: dict[LocationBucket, float] = {
    LocationBucket.REMOTE: 1.0,
    LocationBucket.VANCOUVER: 0.85,
    LocationBucket.TORONTO: 0.7,
    LocationBucket.OTHER_CANADA: 0.4,
    LocationBucket.OTHER: 0.0,
}

# Recency given to a posting whose date couldn't be parsed (LLD §6.3). It sits
# below a fresh dated posting but above a near-stale one, so date_unknown jobs are
# still surfaced and ranked rather than silently lost (spec §7).
_DATE_UNKNOWN_RECENCY = 0.3  # LLD §6.3


def _clamp01(value: float) -> float:
    """Clamp a float to the closed unit interval [0, 1]."""
    return max(0.0, min(1.0, value))


def _cosine(a: NDArray[np.float32], b: NDArray[np.float32]) -> float:
    """Cosine similarity of two vectors; 0.0 if either has zero magnitude."""
    denom = float(np.linalg.norm(a)) * float(np.linalg.norm(b))
    if denom == 0.0:
        return 0.0
    return float(np.dot(a, b) / denom)


def matched_skills(text: str, must_have_skills: list[str]) -> list[str]:
    """Return the must-have skills present in ``text`` (LLD §6.3 / §9.2).

    Matches are word-boundary and case-insensitive, so ``java`` matches "Java"
    but not "javascript". Shared by the scorer's skill component and the
    dashboard's matched-skill chips so the matching logic lives in one place.
    """
    return [
        skill
        for skill in must_have_skills
        if re.search(rf"\b{re.escape(skill)}\b", text, flags=re.IGNORECASE)
    ]


def _skill_score(text: str, must_have_skills: list[str]) -> float:
    """Fraction of must-have skills present in ``text`` (LLD §6.3).

    Saturates at 1.0 once every must-have skill is present.
    """
    needed = len(must_have_skills)
    if needed == 0:
        return 0.0
    return min(1.0, len(matched_skills(text, must_have_skills)) / needed)


def _recency_score(job: Job, *, max_age_days: int, now: datetime) -> float:
    """Linear recency decay over the 0..max_age_days window (LLD §6.3).

    Newest ≈ 1.0, a posting at the cutoff ≈ 0.0; ``date_unknown`` postings get a
    fixed mid-low score so they still rank.
    """
    if job.date_unknown or job.posted_at is None:
        return _DATE_UNKNOWN_RECENCY
    age_days = (now - job.posted_at).days
    return _clamp01(1.0 - age_days / max_age_days)


def score_job(
    job: Job,
    profile_vec: NDArray[np.float32],
    job_vec: NDArray[np.float32],
    *,
    profile: Profile,
    weights: Weights,
    now: datetime,
) -> ScoreBreakdown:
    """Compute the weighted match score and its component breakdown (LLD §6.3–§6.4).

    ``profile_vec``/``job_vec`` are the L2-normalized embeddings from
    :func:`build_profile_vector`/:func:`embed_job`. They are passed in rather than
    re-embedded here so this function is pure and model-free — which makes the
    load-bearing ranking test (tasks.md T16) deterministic and offline. (The §8
    pipeline pseudocode abbreviates embedding and scoring into one call; they are
    separated here for that testability.)

    The four components (semantic cosine clamped to [0,1], skill fraction,
    location bonus, recency decay) are combined as the weight-normalized sum from
    LLD §6.4 and scaled to a 0–100 ``final``.
    """
    semantic = _clamp01(_cosine(profile_vec, job_vec))
    skill = _skill_score(f"{job.title}\n{job.description}", profile.must_have_skills)
    location = _LOCATION_BONUS[job.location_bucket]
    recency = _recency_score(job, max_age_days=profile.max_age_days, now=now)

    # Denominator is guaranteed > 0 by the Weights validator (settings.py).
    denom = weights.semantic + weights.skill + weights.location + weights.recency
    final01 = (
        weights.semantic * semantic
        + weights.skill * skill
        + weights.location * location
        + weights.recency * recency
    ) / denom
    return ScoreBreakdown(
        final=round(100.0 * final01, 1),
        semantic=semantic,
        skill=skill,
        location=location,
        recency=recency,
        scored_at=now,
    )


__all__ = [
    "Encoder",
    "extract_resume",
    "load_model",
    "render_targeting",
    "build_profile_vector",
    "embed_job",
    "score_job",
    "matched_skills",
]
